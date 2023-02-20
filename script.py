import os
from docx import Document
import sys


def main(accumAllSpecs):
    name_of_script = sys.argv[0]
    folder_name = sys.argv[1]
    new_file_name = sys.argv[2]
    accum_dict = accumAllSpecs(folder_name)
    createSpecDocFileFromDict(accum_dict, new_file)

def getSpecificationList(mypath):
    onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
    filtered = filter(lambda spec: ((spec.contains("Cпецификация")
                                     or spec.contains("СО"), onlyfiles)))
    return filtered

def getMatrixFromFile(filename):
    doc = Document("filename")
    tables = doc.tables[0]
    ls =[]
    for row in tables.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                ls.append(paragraph.text)

def createDictFromMatrix(matrix):
    new_dict = {}
    first_row = matrix[0]
    name_index = first_row.index("Наименование и техническая спецификация")
    amount_index = first_row.index("Кол-во")
    type_mesurement_ind = first_row.index("Ед. изм.")
    for rows in matrix:
        new_dict = new_dict + {first_row[name_index]:(first_row[amount_index],
                                                      first_row[type_mesurement_ind])}
    return new_dict

def accumAllSpecs(path):
    accum_dict = {}
    file_list = getSpecificationList(path)
    for file_ in file_list:
        spec_matrix = getMatrixFromFile(file_)
        spec_dict = createDictFromMatrix(spec_matrix)
        for key in spec_dict.keys():
                accum_dict[k] = accum_dict.get(k, 0) + spec_dict[k]
    return accum_dict
def createSpecDocFileFromDict(accum_dict, new_file):
    doc=Document()
    doc.add_heading('Accum specification', 0)
    table = doc.add_table(rows = len(accum_dict) + 1, 3)
    header  = table.rows[0]
    header[0] = "Наименование и техническая спецификация"
    header[1] = "Eд. изм."
    header[2] = "Кол-во"
    table_index = 1
    for key, vals in accum_dict.items():
        row = table.rows[table_index]
        row[0] = kay
        row[1] = vals[0]
        row[2] = vals[1]
        table_index+=1
    doc.add_page_break()
    document.save(f'{new_file}.docx')


if __name__ == "__main__":
    main()

